import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Configuration du style par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    if not os.path.exists(md_path):
        print(f"Erreur : Le fichier {md_path} n'existe pas.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip('\n')
        
        # Gestion des tableaux Markdown
        if '|' in line and '---' not in line:
            in_table = True
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                table_data.append(cells)
            continue
        elif in_table and ('---' in line or not line.strip()):
            if not line.strip() and table_data:
                # Fin du tableau, on l'écrit
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        table.cell(i, j).text = cell_data
                table_data = []
                in_table = False
            continue
        elif in_table and not '|' in line:
             in_table = False
             table_data = []

        # Titres
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
        elif line.startswith('---'):
            doc.add_page_break()
        elif line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.style = 'Quote'
        elif line.strip():
            # Texte normal avec gestion du gras simples
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph()

    # Si un tableau était en cours à la fin du fichier
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_data):
            for j, cell_data in enumerate(row_data):
                table.cell(i, j).text = cell_data

    doc.save(docx_path)
    print(f"Succès : {docx_path} généré.")

if __name__ == "__main__":
    # Définir les chemins
    base_dir = r"C:\Users\LUXE\.gemini\antigravity\brain\b341945e-f7b3-4b03-a927-51136d2faec0"
    output_dir = r"c:\Users\LUXE\Desktop\tender-analyzer"
    
    files_to_convert = [
        ("presentation_commerciale_unifiee.md", "Presentation_Commerciale_AlertesPME.docx"),
        ("analyse_roi_tarification.md", "Analyse_ROI_AlertesPME.docx")
    ]
    
    for md_name, docx_name in files_to_convert:
        md_p = os.path.join(base_dir, md_name)
        docx_p = os.path.join(output_dir, docx_name)
        convert_md_to_docx(md_p, docx_p)
